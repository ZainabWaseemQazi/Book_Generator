from fastapi import APIRouter
from fastapi.responses import HTMLResponse
from docx import Document

from app.database import supabase

router = APIRouter()


@router.get("/compile-book/{book_id}")
def compile_book(book_id: str):

    book = supabase.table("books").select("*").eq("id", book_id).execute().data[0]

    chapters = supabase.table("chapters") \
        .select("*") \
        .eq("book_id", book_id) \
        .order("chapter_number") \
        .execute().data

    document = Document()
    document.add_heading(book["title"], level=1)

    for chapter in chapters:
        document.add_heading(chapter["chapter_title"], level=2)
        document.add_paragraph(chapter["chapter_content"])

    file_path = f"outputs/books/{book_id}.docx"

    document.save(file_path)

    return HTMLResponse(
        f"""
        <h1>Book Generated Successfully</h1>
        <a href='/{file_path}'>Download Book</a>
        """
    )